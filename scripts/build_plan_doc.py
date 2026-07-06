from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "doc" / "biscuitLab_dev_plan.docx"


def set_east_asia_font(style, font_name: str) -> None:
    style.font.name = font_name
    style._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.width = Inches(widths[i])
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, "F2F4F7")
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
            cells[i].width = Inches(widths[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    doc.add_paragraph()


def main() -> None:
    OUT.parent.mkdir(exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    set_east_asia_font(normal, "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = styles[name]
        set_east_asia_font(style, "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    footer = section.footer.paragraphs[0]
    footer.text = "biscuitLab 开发计划"
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.runs[0].font.size = Pt(9)
    footer.runs[0].font.color.rgb = RGBColor(100, 116, 139)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    title_run = title.add_run("biscuitLab 开发计划")
    title_run.font.name = "Microsoft YaHei"
    title_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    title_run.font.size = Pt(26)
    title_run.font.color.rgb = RGBColor(11, 37, 69)
    title_run.bold = True

    subtitle = doc.add_paragraph("基于 TangLab_Codex 开发文档整理 / 前端优先实现版")
    subtitle.runs[0].font.color.rgb = RGBColor(85, 85, 85)
    subtitle.paragraph_format.space_after = Pt(14)

    lead = doc.add_paragraph()
    lead.add_run("项目定位：").bold = True
    lead.add_run(
        "biscuitLab 是一个暗黑科技风个人创意内容站，第一阶段先把前台视觉和主要浏览路径做出来，"
        "后端保留可运行、可扩展的 FastAPI 基础壳。"
    )

    doc.add_heading("1. MVP 范围", level=1)
    add_bullets(
        doc,
        [
            "前台先完成首页、文章列表、文章详情、图片资源库、作品集、Lab 占位页，全部使用 mock 数据。",
            "首页作为第一视觉入口，重点完成暗黑背景、点阵网格、霓虹渐变、玻璃拟态卡片和 hover 动效。",
            "后端先完成 FastAPI 可运行壳：配置、CORS、健康检查、基础 API 占位路由。",
            "后台管理系统暂不展开实现，第一阶段只在计划中保留登录、文章、图片、作品管理边界。",
        ],
    )

    doc.add_heading("2. 技术拆分", level=1)
    add_table(
        doc,
        ["模块", "当前策略", "后续扩展"],
        [
            ["frontend", "Nuxt3 / Vue3 / TypeScript / Tailwind CSS，先做完整前台 mock 体验。", "接入 FastAPI；补 loading、empty、error 状态；加强 SEO。"],
            ["backend", "FastAPI 壳，提供 /health 和文章、图片、作品、登录占位接口。", "接 MySQL、SQLAlchemy、JWT、上传、本地 uploads。"],
            ["admin", "本阶段不抢实现，避免拖慢前台视觉定型。", "Vue3 / Element Plus / Pinia / Axios，做内容管理后台。"],
            ["doc", "沉淀开发计划、启动说明、接口约定。", "后续补部署手册和联调记录。"],
        ],
        [1.15, 2.75, 2.6],
    )

    doc.add_heading("3. 开发阶段计划", level=1)
    add_table(
        doc,
        ["阶段", "目标", "验收标准"],
        [
            ["阶段 1：前台视觉", "完成首页、文章、图片库、作品集和 Lab 占位。", "PC 和移动端不崩版；首页有明显视觉冲击；主要导航可用。"],
            ["阶段 2：后端基础", "补齐模型、schemas、统一响应、JWT 登录和数据库初始化。", "FastAPI 可启动；接口文档可访问；认证边界清晰。"],
            ["阶段 3：后台管理", "实现登录、文章管理、图片管理、作品管理。", "能发布文章、上传图片、维护作品；后台风格简洁稳定。"],
            ["阶段 4：前后端联调", "前台从 mock 数据切到真实接口。", "列表、详情、图片下载/复制链接、作品展示均走 API。"],
            ["阶段 5：上线与二期", "Docker/Nginx/环境变量/部署文档，之后再做留言板和 Lab 实验。", "本地和服务器均可运行；二期范围不挤压 MVP。"],
        ],
        [1.35, 2.45, 2.7],
    )

    doc.add_heading("4. 前端页面规划", level=1)
    add_bullets(
        doc,
        [
            "首页 /：biscuitLab 品牌 Hero、精选作品、最新文章、图片资源预览、联系方式。",
            "文章 /articles 与 /articles/:slug：搜索、分类、标签、卡片展示、详情内容、目录和上一篇/下一篇。",
            "图片库 /gallery：瀑布流、分类筛选、大图预览、下载、复制链接、下载次数展示。",
            "作品集 /works：项目卡片、技术栈、精选标记、在线预览和 GitHub 链接。",
            "Lab /lab：第二阶段实验入口，预留粒子背景、AI 对话流、3D 卡片等方向。",
        ],
    )

    doc.add_heading("5. 后端壳规划", level=1)
    add_bullets(
        doc,
        [
            "健康检查：GET /health。",
            "前台占位接口：GET /api/articles、GET /api/images、GET /api/works。",
            "登录占位接口：POST /api/auth/login。",
            "后续补齐 users、articles、categories、tags、images、works、messages 数据模型。",
            "后续引入 MySQL、SQLAlchemy、Pydantic schemas、JWT、上传目录、统一异常处理。",
        ],
    )

    doc.add_heading("6. 当前交付清单", level=1)
    add_bullets(
        doc,
        [
            "已创建 frontend Nuxt3 项目骨架和暗黑科技风页面。",
            "已创建 backend FastAPI 最小服务骨架。",
            "已将项目命名调整为 biscuitLab。",
            "已将原始文档内容整理为本开发计划，存放在 doc 目录。",
        ],
    )

    doc.add_heading("7. 下一步建议", level=1)
    add_bullets(
        doc,
        [
            "优先人工查看首页视觉，确定是否要更偏霓虹、极客、极简或作品集风。",
            "确认前台 mock 内容字段后，再开始后端数据模型和接口实现。",
            "图片库是差异化功能，下载、复制链接、大图预览应作为第一版重点打磨。",
            "后台保持效率优先，不使用过重动效，避免和前台视觉目标混淆。",
        ],
    )

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
